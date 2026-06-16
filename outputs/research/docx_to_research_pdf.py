from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    PageBreak,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


DOCX = Path("outputs/research/persona_simulation_research_paper_ko.docx")
PDF = Path("outputs/research/persona_simulation_research_paper_ko.pdf")


def register_font():
    candidates = [
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/NanumGothic.ttf"),
    ]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("KoreanBody", str(path)))
            return "KoreanBody"
    return "Helvetica"


FONT = register_font()


def clean(text):
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def make_styles():
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "KRBody",
        parent=styles["BodyText"],
        fontName=FONT,
        fontSize=9.8,
        leading=14.2,
        textColor=colors.HexColor("#141414"),
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    return {
        "body": base,
        "title": ParagraphStyle(
            "KRTitle",
            parent=base,
            fontSize=18,
            leading=24,
            textColor=colors.HexColor("#1F3A5F"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "KRSubtitle",
            parent=base,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#5A5A5A"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "h1": ParagraphStyle(
            "KRH1",
            parent=base,
            fontSize=14.5,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "KRH2",
            parent=base,
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=9,
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "KRSmall",
            parent=base,
            fontSize=8.5,
            leading=11.5,
            textColor=colors.HexColor("#333333"),
            spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "KRBullet",
            parent=base,
            leftIndent=16,
            firstLineIndent=-10,
            spaceAfter=4,
        ),
    }


STYLES = make_styles()


class PaperDoc(BaseDocTemplate):
    def __init__(self, filename):
        super().__init__(
            filename,
            pagesize=letter,
            leftMargin=1 * inch,
            rightMargin=1 * inch,
            topMargin=0.85 * inch,
            bottomMargin=0.75 * inch,
        )
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="normal",
        )
        self.addPageTemplates([PageTemplate(id="paper", frames=[frame], onPage=self.decorate)])

    def decorate(self, canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT, 8.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(1 * inch, 10.55 * inch, "Persona Collection Lab 연구 논문 초안")
        canvas.drawRightString(7.5 * inch, 0.45 * inch, str(doc.page))
        canvas.restoreState()


def para_style(paragraph, idx):
    name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    if idx == 0:
        return STYLES["title"]
    if idx in (1, 2):
        return STYLES["subtitle"]
    if name == "Heading 1":
        return STYLES["h1"]
    if name in ("Heading 2", "Heading 3"):
        return STYLES["h2"]
    if name.startswith("List Bullet"):
        return STYLES["bullet"]
    if name.startswith("List Number"):
        return STYLES["bullet"]
    if text.startswith("["):
        return STYLES["small"]
    return STYLES["body"]


def iter_blocks(doc):
    body = doc._element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            for p in doc.paragraphs:
                if p._p is child:
                    yield ("p", p)
                    break
        elif child.tag.endswith("}tbl"):
            for t in doc.tables:
                if t._tbl is child:
                    yield ("tbl", t)
                    break


def build_pdf():
    source = Document(DOCX)
    story = []
    paragraph_index = 0
    for kind, obj in iter_blocks(source):
        if kind == "p":
            text = obj.text.strip()
            if not text:
                continue
            style = para_style(obj, paragraph_index)
            paragraph_index += 1
            prefix = ""
            if obj.style and obj.style.name.startswith("List Bullet"):
                prefix = "- "
            elif obj.style and obj.style.name.startswith("List Number"):
                prefix = ""
            story.append(Paragraph(clean(prefix + text), style))
        else:
            rows = []
            for row in obj.rows:
                rows.append([
                    Paragraph(clean(cell.text.strip()), STYLES["small"]) for cell in row.cells
                ])
            if rows:
                col_count = len(rows[0])
                usable = 6.5 * inch
                col_widths = [usable / col_count] * col_count
                table = Table(rows, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
                table.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), FONT),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D0DA")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAB4C2")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.extend([Spacer(1, 4), table, Spacer(1, 8)])
    PaperDoc(str(PDF)).build(story)
    reader = PdfReader(str(PDF))
    print(f"{PDF} pages={len(reader.pages)}")


if __name__ == "__main__":
    build_pdf()
