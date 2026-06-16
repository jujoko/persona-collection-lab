from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "outputs/research/persona_simulation_research_paper_ko.docx"


NAVY = RGBColor(31, 58, 95)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(20, 20, 20)
GRAY = RGBColor(90, 90, 90)
FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"


def set_run_font(run, name="Malgun Gothic", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.header.paragraphs[0].text = "Persona Collection Lab 연구 논문 초안"
    for run in section.header.paragraphs[0].runs:
        set_run_font(run, size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Malgun Gothic"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_para(doc, text="", style=None, bold=False, italic=False, size=None, color=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.4)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(text)
            set_run_font(r, size=9.2)
    table_geometry(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_document(doc)

    add_para(doc, "게임 기반 캐릭터 시뮬레이션을 통한 LLM 페르소나 표현공간 학습 프레임워크", size=20, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "Persona Collection Lab 연구 논문 초안", size=12.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "작성일: 2026-06-17 | 대상: 게임 개발 및 페르소나 연구 기획", size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    add_para(doc, "초록", style="Heading 1")
    abstract = (
        "본 논문은 사용자가 자유롭게 작성한 캐릭터 프롬프트와 생애 사건 시뮬레이션 로그를 결합하여, "
        "LLM 기반 페르소나 표현공간을 학습하기 위한 게임형 데이터 수집 프레임워크를 제안한다. "
        "기존 페르소나 에이전트 연구는 대체로 성격, 배경, 말투, 관계와 같은 사람이 미리 정의한 속성 목록을 "
        "모델 입력으로 제공하고, 응답의 일관성을 평가하는 방식에 머무르는 경우가 많다. 그러나 실제 페르소나는 "
        "고정된 속성표라기보다 상황, 선택, 결과, 피드백을 거치며 드러나는 행동 경향의 네트워크에 가깝다. "
        "본 연구는 이를 위해 M1(프롬프트-잠재상태 인코더), M2(잠재상태 기반 행동 결정 모델), "
        "M3(잠재 표현공간 구조 분석 모델)를 분리하고, 생애 arc와 story flag를 통해 해석 가능한 경로를 "
        "보존하는 설계를 제안한다. 또한 플레이어를 직접 조종자가 아니라 '자신이 만든 캐릭터의 행동을 예측하고 "
        "검증하는 관찰자'로 배치함으로써, 게임적 재미와 연구 데이터 품질을 동시에 확보하는 방법을 논의한다."
    )
    add_para(doc, abstract)
    add_para(doc, "주요어: LLM 페르소나, 캐릭터 시뮬레이션, 행동 기반 표현학습, 인터랙티브 스토리, 게임 기반 데이터 수집", italic=True, color=GRAY)

    add_para(doc, "1. 서론", style="Heading 1")
    for text in [
        "LLM 기반 캐릭터 또는 역할 수행 에이전트는 사용자가 제공한 설명문을 바탕으로 특정 인물처럼 말하고 행동하도록 설계된다. 이 접근은 구현이 쉽고 사용자가 즉각적으로 캐릭터성을 확인할 수 있다는 장점이 있다. 그러나 연구 관점에서 보면 중요한 한계가 있다. 사용자가 제공한 문장과 연구자가 미리 정한 성격 분류가 곧 페르소나의 전부가 되어 버리기 때문이다.",
        "본 연구의 출발점은 다르다. 우리는 페르소나를 '정의된 성격표'가 아니라 '반복되는 상황 판단과 행동 결과로부터 추정되는 잠재 구조'로 본다. 예를 들어 어떤 인물이 겁이 많지만 관계를 버리지 못하는지, 성취 욕구가 강하지만 불공정 앞에서 멈춰서는지, 권위에 순응하지만 특정 관계에서는 저항하는지는 단일 성격 축으로 잘 설명되지 않는다. 이런 차이는 사건 속에서 어떤 선택을 하는가, 그 선택이 이후 사건과 자기 이해를 어떻게 바꾸는가를 통해 드러난다.",
        "Persona Collection Lab은 이 문제를 게임으로 다룬다. 사용자는 자유로운 문장으로 캐릭터를 만들고, 시스템은 그 문장을 잠재 벡터로 변환한다. 이후 캐릭터는 성장기와 생애 사건을 지나며 선택과 결과를 누적한다. 플레이어는 캐릭터를 조종하기보다는 '이 캐릭터라면 무엇을 할까'를 예측하고, 결과가 캐릭터답다고 느껴지는지 피드백한다. 이 과정은 자연스럽게 연구 데이터가 된다."
    ]:
        add_para(doc, text)

    add_callout(doc, "핵심 주장", "페르소나는 사용자가 처음 입력한 설명문만으로 고정되지 않는다. 프롬프트, 사건, 행동, 결과, 피드백의 연쇄를 통해 점진적으로 관찰되고 학습되는 잠재 표현이어야 한다.")

    add_para(doc, "2. 관련 연구", style="Heading 1")
    add_para(doc, "2.1 LLM 역할 수행과 페르소나 일관성", style="Heading 2")
    for text in [
        "CharacterGPT는 소설 요약으로부터 인물의 성격, 배경, 관계를 재구성하여 역할 수행 에이전트의 페르소나 일관성을 높이려는 연구이다[1]. 이는 문서 기반 페르소나가 그대로 입력될 때 정보 추출이 흔들리고 중요한 배경이 누락될 수 있음을 지적한다. 본 연구는 이 문제의식을 공유하지만, 인물 정보를 재구성하는 데서 한 걸음 더 나아가 사건 속 행동 로그를 페르소나 학습의 핵심 신호로 삼는다.",
        "대화형 캐릭터 연구는 대개 '그 인물답게 말하는가'를 주요 평가 대상으로 삼는다. 반면 Persona Collection Lab은 '그 인물답게 판단하고 변하는가'를 묻는다. 이 차이는 데이터 구조에도 영향을 준다. 발화 텍스트만 저장하는 것이 아니라 사건, 선택지, 예측 행동, 실제 행동, latent before/after, story flag before/after를 함께 저장해야 한다."
    ]:
        add_para(doc, text)

    add_para(doc, "2.2 LLM 기반 인터랙티브 스토리와 게임 생성", style="Heading 2")
    for text in [
        "STORY2GAME은 LLM이 스토리, 세계, 행동 코드를 생성하여 인터랙티브 픽션을 구성하는 접근을 제안한다[2]. 이 연구의 중요한 시사점은 행동의 precondition과 effect가 게임 상태 추적의 핵심이라는 점이다. 본 연구의 story flag와 action sets 역시 같은 문제를 다룬다. 다만 본 연구는 열린 행동 생성보다 연구 가능한 선택 로그를 우선하므로, 고정된 사건 구조와 제한된 선택지를 유지하면서 선택지 노출과 결과 문장을 다양화한다.",
        "Word2World와 interactive fiction world generation 연구 역시 LLM이 이야기와 세계 구조를 생성하는 가능성을 보여준다[3][4]. 그러나 본 연구는 스토리 생성 자체보다 페르소나 표현 학습을 목적으로 한다. 따라서 사건 구조는 완전히 자유롭게 생성하기보다 비교 가능한 단위로 통제되어야 하며, AI는 표면 문장 변주와 해석 보조에 더 적합하다."
    ]:
        add_para(doc, text)

    add_para(doc, "2.3 게임화와 동기 설계", style="Heading 2")
    for text in [
        "게임화 연구는 점수, 배지, 리더보드 같은 요소만으로 동기가 자동 발생하지 않으며, 자율성, 유능감, 관계성 같은 심리적 욕구와 결합될 때 효과가 커진다고 본다[5][6]. 본 시스템은 플레이어에게 긴 설문을 요구하지 않고, 캐릭터 생성, 예측, 결과 확인, 피드백, 재시뮬레이션이라는 짧은 루프를 제공한다. 이는 데이터 수집이 플레이 경험 안에 자연스럽게 들어가야 한다는 설계 원칙과 맞닿아 있다.",
        "중요한 점은 플레이어가 캐릭터를 직접 선택으로 조종하지 않는다는 것이다. 조종 게임으로 만들면 수집되는 데이터는 플레이어의 선택 데이터가 된다. 본 연구가 원하는 것은 캐릭터가 스스로 그럴듯하게 행동하는지에 대한 예측과 평가 데이터다. 따라서 사용자 개입은 캐릭터 생성, 행동 예측, 피드백, 재실험에 집중된다."
    ]:
        add_para(doc, text)

    add_para(doc, "3. 연구 문제", style="Heading 1")
    add_para(doc, "본 연구는 다음 네 가지 질문을 중심으로 한다.")
    for item in [
        "RQ1. 자유 형식 캐릭터 프롬프트는 어떤 초기 latent persona 상태로 변환될 수 있는가?",
        "RQ2. 같은 캐릭터가 생애 사건을 거치며 선택과 결과에 따라 어떻게 변화하는가?",
        "RQ3. 플레이어는 자신이 만든 캐릭터의 행동을 얼마나 잘 예측하며, 그 불일치는 어떤 학습 신호가 되는가?",
        "RQ4. story flag처럼 해석 가능한 생애 이력과 latent vector처럼 연속적인 잠재 상태를 어떻게 함께 사용할 수 있는가?"
    ]:
        add_number(doc, item)

    add_para(doc, "4. 시스템 설계", style="Heading 1")
    add_para(doc, "4.1 전체 구조", style="Heading 2")
    add_para(doc, "제안 시스템은 세 모델과 하나의 스토리 런타임으로 구성된다. M1은 사용자 프롬프트를 latent seed로 변환한다. M2는 현재 latent, 사건, 선택지 정보를 바탕으로 캐릭터가 고를 행동을 결정한다. M3는 수집된 로그를 바탕으로 latent 표현공간의 차원, 연결, 업데이트 규칙을 분석한다. Story runtime은 해석 가능한 생애 흐름을 유지하기 위해 arc, event, action, flag를 관리한다.")
    add_matrix(
        doc,
        ["구성요소", "입력", "출력", "연구적 의미"],
        [
            ["M1", "자유 프롬프트, 이름, seed", "초기 latent, 근거 조각", "사용자 언어가 잠재 표현으로 투사되는 방식"],
            ["M2", "latent, event, visible actions", "행동 ID, 점수, 해석", "상황별 행동 예측 신호"],
            ["M3", "행동 로그, 피드백, latent 변화", "schema, edge, sensitivity", "페르소나 표현공간의 구조 발견"],
            ["Story Runtime", "arc, flags, action sets", "다음 사건, 선택지 노출", "비교 가능한 생애 맥락 제공"],
        ],
        [1700, 2450, 2300, 2910],
    )

    add_para(doc, "4.2 latent와 story flag의 분리", style="Heading 2")
    for text in [
        "본 시스템에서 latent vector와 story flag는 서로 대체 관계가 아니다. latent는 캐릭터의 부드러운 행동 경향을 표현하며, 연구자가 미리 이름 붙인 성격 축에 고정되지 않는다. 반면 story flag는 수능 결과, 전공, 관계 이력, 평판, 건강 상태처럼 이후 사건 분기를 해석 가능하게 만드는 생애 이력이다.",
        "이 분리는 연구적으로 중요하다. story flag만 있으면 캐릭터는 정해진 루트의 조합이 되고, latent만 있으면 왜 어떤 사건이 열렸는지 설명하기 어렵다. 두 층을 함께 쓰면 '왜 이 선택지가 보였는가'와 '왜 이 캐릭터가 그 선택을 했는가'를 따로 분석할 수 있다."
    ]:
        add_para(doc, text)

    add_para(doc, "4.3 생애 arc와 사건 구조", style="Heading 2")
    add_para(doc, "현재 콘텐츠 설계는 초등학교, 중학교, 고등학교, 재수, 대학교, 취업 준비, 사회초년생, 30대, 중년으로 이어지는 생애 arc를 가진다. 각 arc에는 4~10개의 사건 후보가 있고, 한 플레이에서는 조건과 seed에 따라 일부만 노출된다. 사건은 2~4개의 action을 가지며, action-level requires/excludes를 통해 같은 사건에서도 보이는 선택지가 달라질 수 있다.")
    add_matrix(
        doc,
        ["Arc", "주요 역할", "대표 flags"],
        [
            ["Arc 0~1", "어린 시절 반응 패턴과 초기 관심사 형성", "childhood_pattern, first_interest, social_role"],
            ["Arc 2~3", "시험, 전공, 대학, 첫 독립 경험", "exam_score, major, university_tier"],
            ["Arc 4~5", "취업과 조직 내 첫 평판", "career_start, reputation, financial_status"],
            ["Arc 6~7", "가족, 건강, 회고, 엔딩 표면", "family_structure, health, accumulated flags"],
        ],
        [1500, 4800, 3060],
    )

    add_para(doc, "5. 게임 루프 설계", style="Heading 1")
    add_para(doc, "플레이어는 캐릭터를 직접 조종하지 않는다. 대신 자신이 만든 캐릭터가 어떤 행동을 할지 예측한다. 이 설계는 게임적으로는 추리와 수집의 재미를 만들고, 연구적으로는 캐릭터 일관성에 대한 인간 평가 데이터를 만든다.")
    for item in [
        "캐릭터 생성: 사용자는 이름과 자유 프롬프트를 입력한다.",
        "초기 해석: M1이 prompt-conditioned latent seed와 근거 조각을 생성한다.",
        "사건 제시: story runtime이 현재 flags와 seed에 맞는 사건과 선택지를 제시한다.",
        "예측: 플레이어는 캐릭터가 고를 것 같은 선택지를 고른다.",
        "행동 결정: M2가 latent와 선택지를 바탕으로 실제 행동을 결정한다.",
        "결과 공개: 결과 문장, 예측 일치 여부, 변화 요약을 보여준다.",
        "피드백: 플레이어는 결과가 캐릭터답다고 느끼는지 가볍게 평가한다.",
        "기록: event log에 latent, flags, 예측, M2 선택, 피드백을 저장한다."
    ]:
        add_number(doc, item)

    add_para(doc, "6. 데이터 스키마", style="Heading 1")
    add_para(doc, "데이터는 연구의 중심 산출물이다. 한 이벤트 단위 로그는 다음 정보를 포함해야 한다.")
    add_matrix(
        doc,
        ["필드 묶음", "예시 필드", "용도"],
        [
            ["식별자", "run_id, character_id, event_id, arc_id", "반복 실험과 경로 재구성"],
            ["입력", "prompt, generated_name, seed", "초기 조건 분석"],
            ["선택 구조", "visible_action_ids, player_predicted_action_id, m2_selected_action_id", "예측 정확도와 선택지 노출 분석"],
            ["상태 변화", "latent_before/after, flags_before/after", "페르소나 변화와 스토리 이력 분석"],
            ["평가", "feedback_kind, consistency_score, fun_score", "선호학습과 품질 평가"],
        ],
        [1800, 3850, 3710],
    )
    add_para(doc, "이 스키마의 장점은 사후 분석이 쉽다는 점이다. 예를 들어 특정 event type에서 latent 변화량이 큰지, 특정 flag 조합이 어떤 행동을 많이 여는지, 플레이어 예측과 M2 선택이 불일치할 때 어떤 피드백이 붙는지를 직접 계산할 수 있다.")

    add_para(doc, "7. 평가 계획", style="Heading 1")
    add_para(doc, "본 연구의 평가는 모델 정확도 하나로 끝나지 않는다. 게임성과 연구 품질을 동시에 봐야 한다.")
    add_matrix(
        doc,
        ["평가 축", "측정 방법", "해석"],
        [
            ["행동 일관성", "플레이어가 '캐릭터답다'고 평가한 비율", "M2가 프롬프트와 상황을 잘 반영하는지"],
            ["예측 가능성", "player prediction과 M2 action의 일치율", "캐릭터가 이해 가능한 패턴을 갖는지"],
            ["경로 다양성", "동일 프롬프트 seed별 unique route 수", "반복 플레이 가치와 데이터 다양성"],
            ["latent 민감도", "event별 latent shift 평균", "어떤 사건이 페르소나를 많이 바꾸는지"],
            ["재방문 의향", "want_to_retry, ending collection", "게임 루프의 지속 가능성"],
        ],
        [1900, 3800, 3660],
    )

    add_para(doc, "8. 현재 프로토타입 상태와 위험", style="Heading 1")
    for text in [
        "현재 프로토타입은 M1 latent seed, M2 rule-based decision, Supabase 저장, story content validator, story adapter의 초기 버전을 갖추고 있다. 또한 arc0부터 arc7까지 생애 사건 초안이 작성되어 있다. 다만 전체 arc가 런타임과 검증에 완전히 편입된 상태는 아니며, 일부 테스트는 arc0과 arc2만 대상으로 한다.",
        "가장 큰 기술적 위험은 전체 run의 사건을 처음부터 한 번에 샘플링하는 구조다. 이 구조에서는 action sets가 이후 사건 후보에 즉시 반영되지 않는다. 따라서 다음 개발 단계에서는 buildRunEvents 중심 구조를 createRun, nextEvent, applyAction 기반 stepper 구조로 전환해야 한다.",
        "두 번째 위험은 Play Model을 너무 빨리 기본 행동 결정기로 승격하는 것이다. 기존 benchmark에서 Play Model은 제한된 정확도와 특정 윤리적 선택 편향을 보였다. 따라서 당분간은 rule-based M2를 기본으로 유지하고, Play Model은 실험 endpoint와 로그 수집 대상으로 두는 편이 안전하다."
    ]:
        add_para(doc, text)

    add_para(doc, "9. 윤리와 개인정보", style="Heading 1")
    for text in [
        "본 시스템은 사용자가 자유롭게 작성한 캐릭터 프롬프트를 저장할 수 있으므로 개인정보와 민감정보 관리가 중요하다. 사용자에게 실제 개인정보를 입력하지 않도록 안내하고, 연구용 export에는 익명화된 character_id와 run_id를 사용해야 한다.",
        "또한 페르소나 추론 결과를 실제 개인의 성격 진단처럼 제시해서는 안 된다. 본 시스템의 latent vector는 연구용 표현이며, 심리검사나 임상적 평가가 아니다. UI에서는 내부 수치를 직접 노출하기보다 '이 캐릭터는 이런 상황에서 이런 경향을 보였다'는 사건 기반 설명으로 제한하는 것이 적절하다.",
        "마지막으로 피드백은 플레이어의 정답 여부를 평가하는 것이 아니라 모델의 캐릭터 일관성에 대한 인간 신호로 해석해야 한다. 플레이어가 틀린 것이 아니라, 시스템이 만든 캐릭터가 예측 가능했는지 혹은 납득 가능했는지를 보는 것이다."
    ]:
        add_para(doc, text)

    add_para(doc, "10. 개발 로드맵", style="Heading 1")
    for item in [
        "모든 arc를 자동 registry로 묶고 content validator 대상에 포함한다.",
        "StoryAdapter를 stepper 구조로 전환하여 선택 결과가 다음 사건 후보에 즉시 반영되게 한다.",
        "플레이어 예측 UI를 추가하고 visible actions, predicted action, M2 selected action을 모두 저장한다.",
        "이벤트 로그에 save_version: 2와 flags_before/after를 추가한다.",
        "엔딩 리포트를 강화해 캐릭터의 주요 변화, 가장 큰 분기, 예측 적중률을 보여준다.",
        "충분한 로그가 쌓이면 M3 schema discovery와 M2 재학습을 수행한다."
    ]:
        add_number(doc, item)

    add_para(doc, "11. 결론", style="Heading 1")
    for text in [
        "본 논문은 자유 프롬프트 기반 캐릭터 생성과 생애 사건 시뮬레이션을 결합하여 LLM 페르소나 표현공간을 학습하는 게임형 연구 프레임워크를 제안했다. 핵심은 페르소나를 사전정의된 성격 축에 고정하지 않고, 사건 속 행동과 피드백을 통해 점진적으로 관찰되는 잠재 구조로 다루는 것이다.",
        "게임적으로는 플레이어가 캐릭터를 예측하고 엔딩과 경로를 수집하는 재미를 얻는다. 연구적으로는 프롬프트, 사건, 선택지, 행동, 결과, 피드백이 정렬된 데이터가 쌓인다. 이 결합은 단순한 성격 테스트도, 단순한 비주얼노벨도 아닌 '페르소나 형성 실험 게임'이라는 독특한 형태를 만든다.",
        "향후 작업은 전체 arc 런타임 통합, stepper 기반 분기, 예측 UI, 연구 로그 스키마 강화, 그리고 누적 데이터 기반 M3/M2 학습으로 이어져야 한다. 이 방향이 유지된다면 Persona Collection Lab은 게임 개발 프로젝트이면서 동시에 LLM 페르소나 연구를 위한 데이터 수집 플랫폼으로 발전할 수 있다."
    ]:
        add_para(doc, text)

    add_para(doc, "참고문헌", style="Heading 1")
    refs = [
        "[1] Park, J., Park, C., & Lim, H. CharacterGPT: A Persona Reconstruction Framework for Role-Playing Agents. arXiv:2405.19778, 2024. https://arxiv.org/abs/2405.19778",
        "[2] Zhou, E., Basavatia, S., Siam, M., Chen, Z., & Riedl, M. O. STORY2GAME: Generating (Almost) Everything in an Interactive Fiction Game. arXiv:2505.03547, 2025. https://arxiv.org/abs/2505.03547",
        "[3] Nasir, M. U., James, S., & Togelius, J. Word2World: Generating Stories and Worlds through Large Language Models. arXiv:2405.06686, 2024. https://arxiv.org/abs/2405.06686",
        "[4] Ammanabrolu, P., Cheung, W., Tu, D., Broniec, W., & Riedl, M. O. Bringing Stories Alive: Generating Interactive Fiction Worlds. arXiv:2001.10161, 2020. https://arxiv.org/abs/2001.10161",
        "[5] Sailer, M., Hense, J. U., Mayr, S. K., & Mandl, H. How gamification motivates: An experimental study of the effects of specific game design elements on psychological need satisfaction. Computers in Human Behavior, 69, 371-380, 2017.",
        "[6] Ryan, R. M., & Deci, E. L. Self-determination theory and the facilitation of intrinsic motivation, social development, and well-being. American Psychologist, 55(1), 68-78, 2000.",
        "[7] Deterding, S., Dixon, D., Khaled, R., & Nacke, L. From game design elements to gamefulness: defining gamification. Proceedings of MindTrek, 2011.",
        "[8] Persona Collection Lab repository documents: STORY_SPEC.md, docs/claude/persona_simulation_research_context.md, docs/codex/play-model-benchmark.md. Accessed 2026-06-17."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(ref)
        set_run_font(r, size=9.4)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
